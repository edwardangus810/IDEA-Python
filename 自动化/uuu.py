# 读取docx中的文本代码示例
import docx


# 获取文档对象
file = docx.Document("one.docx")
print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

# 输出每一段的内容
for para in file.paragraphs:
  print(para.text)

# 输出段落编号及段落内容
for i in range(len(file.paragraphs)):
    print("第" + str(i + 1) + "段的内容是：" + file.paragraphs[i].text)

'''
doc2 = docx.Document()  # 创建一个Document对象
doc2.add_paragraph('time')  # 增加一个paragraph
# 插入有序列表,段落的前面会有序号123
doc2.add_paragraph('把冰箱门打开', style='List Number')
doc2.add_paragraph('把大象装进去', style='List Number')
doc2.add_paragraph('把冰箱门关上', style='List Number')
# 插入无序列表，段落的前面没有序号
doc2.add_paragraph('把冰箱门打开', style='List Bullet')
doc2.add_paragraph('把大象装进去', style='List Bullet')
doc2.add_paragraph('把冰箱门关上', style='List Bullet')
# 插入一个6行6列的表格
table = doc2.add_table(rows=6, cols=6, style='Table Grid')
for i in range(0, 6):
    for j in range(0, 6):
        table.cell(i, j).text = "第{i}行{j}列".format(i=i + 1, j=j + 1)
# 插入照片
doc2.save('two.docx')  # 保存文档
'''