import docx


'''
doc2 = docx.Document()  # 创建一个Document对象
doc2.add_paragraph('time')  # 增加一个paragraph
# 插入有序列表,段落的前面会有序号123
doc2.add_paragraph('把冰箱门打开', style='List Number')
doc2.add_paragraph('把大象装进去', style='List Number')
doc2.add_paragraph('把冰箱门关上', style='List Number')
# 插入无序列表，段落的前面没有序号
doc2.add_paragraph('把冰箱门打开', style='List Bullet')
doc2.add_paragraph('把大象装进去', style='List Bullet')
doc2.add_paragraph('把冰箱门关上', style='List Bullet')
# 插入一个6行6列的表格
table = doc2.add_table(rows=6, cols=6, style='Table Grid')
for i in range(0, 6):
    for j in range(0, 6):
        table.cell(i, j).text = "第{i}行{j}列".format(i=i + 1, j=j + 1)
# 插入照片
doc2.save('two.docx')  # 保存文档
'''


doc3 = docx.Document()
doc3.add_paragraph('time')
# 无序列表
doc3.add_paragraph('今天是星期一', style='List Number')
doc3.add_paragraph('今天是星期二', style='List Number')
doc3.add_paragraph('今天是星期三', style='List Number')
doc3.add_paragraph('今天是星期四', style='List Number')
doc3.add_paragraph('今天是星期五', style='List Number')
doc3.add_paragraph('今天是星期六', style='List Number')
doc3.add_paragraph('今天是星期日', style='List Number')
# 有序列表
doc3.add_paragraph('今天是星期一', style='List Bullet')
doc3.add_paragraph('今天是星期二', style='List Bullet')
doc3.add_paragraph('今天是星期三', style='List Bullet')
doc3.add_paragraph('今天是星期四', style='List Bullet')
doc3.add_paragraph('今天是星期五', style='List Bullet')
doc3.add_paragraph('今天是星期六', style='List Bullet')
doc3.add_paragraph('今天是星期日', style='List Bullet')
# 插入表格
table = doc3.add_table(rows=6, cols=6, style='Table Grid')
for i in range(0, 6):
    for j in range(0, 6):
        table.cell(i, j).text = "第{i}行{j}列".format(i=i + 1, j=j + 1)

doc3.save('three.docx')